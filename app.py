import requests
from docx import Document
import os

endpoint = "ENDPOINT"
subscription_key = "SUA_CHAVE"
location = "brazilsouth"

languageSaida = 'pt-br'

def translate_text(text, language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId' : str(os.urandom(16))
  }

  body = [{
      'text' : text
  }]
  params = {
      'api-version' : '3.0',
      'from' : 'en',
      'to' : language
  }

  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]['translations'][0]['text']

translate_text("Say you're leavin' on a seven-thirty train", "pt-br")

def translate_documente(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translate_text(paragraph.text, languageSaida)
    full_text.append(translated_text)
  
  translated_doc = Document()
  for line in full_text:
    print(line)
    translated_doc.add_paragraph(line)
  path_translated_doc = 'translated_doc.docx'
  translated_doc.save('translated_doc.docx')
  return path_translated_doc

input_file = "Come here.docx"